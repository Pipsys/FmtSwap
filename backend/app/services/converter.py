"""
Document conversion service.
Supported conversion types:
- pdf_to_docx
- pdf_to_jpg (ZIP archive with JPEG pages)
- jpg_to_pdf
- word_to_pdf
"""
import asyncio
import importlib.util
import logging
import os
import re
import shutil
import subprocess
import sys
import tempfile
import textwrap
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional
from zipfile import ZIP_DEFLATED, ZipFile

from sqlalchemy.orm import Session

from app.core.config import get_settings
from app.models.models import ConversionTask, TaskStatus

logger = logging.getLogger(__name__)
settings = get_settings()

PDF_TO_DOCX = "pdf_to_docx"
PDF_TO_JPG = "pdf_to_jpg"
JPG_TO_PDF = "jpg_to_pdf"
WORD_TO_PDF = "word_to_pdf"
HEIC_TO_JPG = "heic_to_jpg"
PNG_TO_WEBP = "png_to_webp"
SVG_TO_PNG = "svg_to_png"
PNG_TO_ICO = "png_to_ico"
PSD_AI_TO_PNG = "psd_ai_to_png"
PSD_AI_TO_JPG = "psd_ai_to_jpg"
VIDEO_TO_MP4 = "video_to_mp4"
GIF_TO_MP4 = "gif_to_mp4"
VIDEO_TO_MP3 = "video_to_mp3"
VIDEO_TO_AVI = "video_to_avi"
VIDEO_TO_MOV = "video_to_mov"
ARCH_ZIP_PACK = "arch_zip_pack"
ARCH_7Z_PACK = "arch_7z_pack"
ARCH_RAR_PACK = "arch_rar_pack"
ARCH_ZIP_UNPACK = "arch_zip_unpack"
ARCH_7Z_UNPACK = "arch_7z_unpack"
ARCH_RAR_UNPACK = "arch_rar_unpack"
COMP_IMG_JPG = "comp_img_jpg"
COMP_IMG_WEBP = "comp_img_webp"
COMP_IMG_AVIF = "comp_img_avif"
COMP_PDF = "comp_pdf"
COMP_OFFICE_ZIP = "comp_office_zip"
COMP_VIDEO_MP4 = "comp_video_mp4"
COMP_AUDIO_MP3 = "comp_audio_mp3"

VIDEO_INPUT_EXTENSIONS = (
    ".mp4",
    ".mov",
    ".avi",
    ".mkv",
    ".webm",
    ".wmv",
    ".flv",
    ".m4v",
    ".mpg",
    ".mpeg",
    ".3gp",
    ".ts",
    ".m2ts",
    ".ogv",
)

AUDIO_INPUT_EXTENSIONS = (
    ".mp3",
    ".wav",
    ".m4a",
    ".aac",
    ".flac",
    ".ogg",
    ".oga",
    ".wma",
)

SCANNED = "scanned"
TEXT = "text"


def _ensure_pymupdf_compat() -> None:
    """
    pdf2docx 0.5.x expects Rect.get_area(), which is missing in newer PyMuPDF.
    Add a compatibility shim at runtime.
    """
    try:
        import fitz  # PyMuPDF
    except Exception:
        return

    def _area(rect) -> float:
        return abs((rect.x1 - rect.x0) * (rect.y1 - rect.y0))

    if hasattr(fitz, "Rect") and not hasattr(fitz.Rect, "get_area"):
        fitz.Rect.get_area = _area
    if hasattr(fitz, "IRect") and not hasattr(fitz.IRect, "get_area"):
        fitz.IRect.get_area = _area


def _ensure_dirs() -> tuple[Path, Path]:
    """Ensure upload and output directories exist. Returns (upload_dir, output_dir)."""
    upload_dir = Path(settings.UPLOAD_DIR)
    output_dir = Path(settings.OUTPUT_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)
    output_dir.mkdir(parents=True, exist_ok=True)
    return upload_dir, output_dir


def _do_convert(pdf_path: str, docx_path: str) -> None:
    """
    Perform the actual PDF->DOCX conversion using pdf2docx.
    This is a blocking call, run in a thread pool.
    """
    _ensure_pymupdf_compat()
    from pdf2docx import Converter

    cv = Converter(pdf_path)
    try:
        # parse() extracts text, images, tables, fonts and writes to DOCX
        cv.convert(docx_path, start=0, end=None)
    finally:
        cv.close()


def _detect_pdf_kind(pdf_path: Path) -> str:
    """
    Heuristic PDF detector:
    - text PDF: at least one sampled page has enough extractable text
    - scanned PDF: mostly image pages without selectable text
    """
    _ensure_pymupdf_compat()
    import fitz  # PyMuPDF

    sample_pages = max(1, settings.PDF_DETECT_SAMPLE_PAGES)
    min_chars_per_page = max(1, settings.PDF_TEXT_MIN_CHARS_PER_PAGE)
    min_text_pages = max(1, settings.PDF_TEXT_MIN_PAGES)

    with fitz.open(str(pdf_path)) as doc:
        if doc.page_count == 0:
            return TEXT

        pages_to_check = min(doc.page_count, sample_pages)
        text_pages = 0
        image_pages = 0
        total_chars = 0

        for page_index in range(pages_to_check):
            page = doc.load_page(page_index)
            page_text = page.get_text("text").strip()
            char_count = len(page_text)
            total_chars += char_count

            if char_count >= min_chars_per_page:
                text_pages += 1
            if page.get_images(full=True):
                image_pages += 1

    logger.info(
        "PDF detection: pages=%s text_pages=%s image_pages=%s total_chars=%s",
        pages_to_check,
        text_pages,
        image_pages,
        total_chars,
    )

    if text_pages >= min_text_pages:
        return TEXT
    if total_chars >= min_chars_per_page * min_text_pages:
        return TEXT
    if image_pages >= max(1, pages_to_check // 2):
        return SCANNED
    return SCANNED


def _resolve_ocrmypdf_command() -> Optional[list[str]]:
    if shutil.which("ocrmypdf"):
        return ["ocrmypdf"]
    if importlib.util.find_spec("ocrmypdf") is not None:
        return [sys.executable, "-m", "ocrmypdf"]
    return None


def _collect_windows_ocr_tool_dirs() -> list[str]:
    """
    Try common Windows install locations for OCR tools when PATH is not configured.
    """
    if os.name != "nt":
        return []

    dirs: list[str] = []

    tesseract_dir = Path(r"C:\Program Files\Tesseract-OCR")
    if (tesseract_dir / "tesseract.exe").exists():
        dirs.append(str(tesseract_dir))

    gs_root = Path(r"C:\Program Files\gs")
    if gs_root.exists():
        for version_dir in sorted((p for p in gs_root.iterdir() if p.is_dir()), reverse=True):
            gs_bin = version_dir / "bin"
            if (gs_bin / "gswin64c.exe").exists() or (gs_bin / "gswin32c.exe").exists():
                dirs.append(str(gs_bin))
                break

    return dirs


def _resolve_ffmpeg_command() -> Optional[str]:
    ffmpeg = shutil.which("ffmpeg")
    if ffmpeg:
        return ffmpeg

    if os.name == "nt":
        candidates = [
            Path(r"C:\ffmpeg\bin\ffmpeg.exe"),
            Path(r"C:\Program Files\ffmpeg\bin\ffmpeg.exe"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)

    return None


def _run_ffmpeg(args: list[str], error_prefix: str) -> None:
    ffmpeg_cmd = _resolve_ffmpeg_command()
    if not ffmpeg_cmd:
        raise RuntimeError(
            "FFmpeg не найден в системе. Установите FFmpeg и добавьте его в PATH."
        )

    process = subprocess.run(
        [ffmpeg_cmd, *args],
        capture_output=True,
        text=True,
    )
    if process.returncode != 0:
        details = (process.stderr or process.stdout or "ошибка FFmpeg").strip()
        raise RuntimeError(f"{error_prefix}: {details}")


def _resolve_imagemagick_command() -> Optional[str]:
    magick = shutil.which("magick")
    if magick:
        return magick

    if os.name == "nt":
        candidates = [
            Path(r"C:\Program Files\ImageMagick-7.1.1-Q16-HDRI\magick.exe"),
            Path(r"C:\Program Files\ImageMagick-7.1.1-Q16\magick.exe"),
            Path(r"C:\Program Files\ImageMagick-7.0.11-Q16-HDRI\magick.exe"),
            Path(r"C:\Program Files\ImageMagick-7.0.11-Q16\magick.exe"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)

    return None


def _run_imagemagick(args: list[str], error_prefix: str) -> None:
    magick_cmd = _resolve_imagemagick_command()
    if not magick_cmd:
        raise RuntimeError(
            "ImageMagick не найден. Установите ImageMagick для конвертации PSD/AI."
        )

    process = subprocess.run(
        [magick_cmd, *args],
        capture_output=True,
        text=True,
    )
    if process.returncode != 0:
        details = (process.stderr or process.stdout or "ошибка ImageMagick").strip()
        raise RuntimeError(f"{error_prefix}: {details}")


def _resolve_ghostscript_command() -> Optional[str]:
    candidates = ["gswin64c", "gswin32c", "gs"]
    for cmd in candidates:
        resolved = shutil.which(cmd)
        if resolved:
            return resolved

    if os.name == "nt":
        gs_root = Path(r"C:\Program Files\gs")
        if gs_root.exists():
            for version_dir in sorted((p for p in gs_root.iterdir() if p.is_dir()), reverse=True):
                gs_bin = version_dir / "bin"
                for exe in ("gswin64c.exe", "gswin32c.exe"):
                    candidate = gs_bin / exe
                    if candidate.exists():
                        return str(candidate)
    return None


def _run_ghostscript(args: list[str], error_prefix: str) -> None:
    gs_cmd = _resolve_ghostscript_command()
    if not gs_cmd:
        raise RuntimeError(
            "Ghostscript не найден в системе. Установите Ghostscript для сжатия PDF."
        )

    process = subprocess.run(
        [gs_cmd, *args],
        capture_output=True,
        text=True,
    )
    if process.returncode != 0:
        details = (process.stderr or process.stdout or "ошибка Ghostscript").strip()
        raise RuntimeError(f"{error_prefix}: {details}")


def _resolve_rar_command() -> Optional[str]:
    rar = shutil.which("rar")
    if rar:
        return rar

    if os.name == "nt":
        candidates = [
            Path(r"C:\Program Files\WinRAR\Rar.exe"),
            Path(r"C:\Program Files (x86)\WinRAR\Rar.exe"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)

    return None


def _unique_name(name: str, used: set[str]) -> str:
    clean = Path(name).name or "file"
    stem = Path(clean).stem
    suffix = Path(clean).suffix
    candidate = clean
    index = 1
    while candidate.lower() in used:
        candidate = f"{stem}_{index}{suffix}"
        index += 1
    used.add(candidate.lower())
    return candidate


def _zip_files(file_items: list[tuple[str, Path]], output_zip_path: Path) -> None:
    used: set[str] = set()
    with ZipFile(output_zip_path, "w", compression=ZIP_DEFLATED, compresslevel=9) as archive:
        for original_name, file_path in file_items:
            arcname = _unique_name(original_name, used)
            archive.write(file_path, arcname=arcname)


def _zip_directory(source_dir: Path, output_zip_path: Path) -> None:
    with ZipFile(output_zip_path, "w", compression=ZIP_DEFLATED, compresslevel=9) as archive:
        for file_path in sorted(source_dir.rglob("*")):
            if file_path.is_file():
                archive.write(file_path, arcname=str(file_path.relative_to(source_dir)))


def _sanitize_member_name(member_name: str) -> Optional[Path]:
    parts = [part for part in Path(member_name.replace("\\", "/")).parts if part not in ("", ".", "..")]
    if not parts:
        return None
    return Path(*parts)


def _build_ocr_env(extra_path_dirs: list[str]) -> dict[str, str]:
    env = os.environ.copy()
    if extra_path_dirs:
        env["PATH"] = os.pathsep.join(extra_path_dirs + [env.get("PATH", "")])
    if settings.TESSDATA_PREFIX:
        env["TESSDATA_PREFIX"] = settings.TESSDATA_PREFIX
    return env


def _read_installed_tesseract_languages(extra_path_dirs: list[str]) -> set[str]:
    env = _build_ocr_env(extra_path_dirs)

    tesseract_cmd = shutil.which("tesseract", path=env.get("PATH"))
    if not tesseract_cmd:
        return set()

    process = subprocess.run(
        [tesseract_cmd, "--list-langs"],
        capture_output=True,
        text=True,
        env=env,
    )
    if process.returncode != 0:
        return set()

    langs: set[str] = set()
    for line in process.stdout.splitlines():
        line = line.strip()
        if not line or line.lower().startswith("list of available languages"):
            continue
        langs.add(line)
    return langs


def _resolve_ocr_lang(extra_path_dirs: list[str]) -> str:
    requested = [x.strip() for x in settings.OCR_LANG.split("+") if x.strip()]
    if not requested:
        requested = ["eng"]

    installed = _read_installed_tesseract_languages(extra_path_dirs)
    if not installed:
        return "+".join(requested)

    available = [lang for lang in requested if lang in installed]
    missing = [lang for lang in requested if lang not in installed]

    if available:
        if missing:
            logger.warning(
                "OCR language fallback: requested=%s, missing=%s, using=%s",
                "+".join(requested),
                "+".join(missing),
                "+".join(available),
            )
        return "+".join(available)

    if "eng" in installed:
        logger.warning(
            "OCR language fallback: requested=%s, available in Tesseract=%s, using=eng",
            "+".join(requested),
            "+".join(sorted(installed)),
        )
        return "eng"

    raise RuntimeError(
        "Отсутствуют языковые данные OCR. Запрошено: "
        f"{'+'.join(requested)}. Установлено: {', '.join(sorted(installed)) or 'нет'}. "
        "Установите языковые пакеты или задайте OCR_LANG с доступным языком."
    )


def _normalize_ocr_text(text: str) -> str:
    """
    Normalize OCR output to reduce common spacing artifacts:
    - remove hyphenation on line breaks
    - convert single line breaks into spaces
    - keep paragraph breaks
    - insert spaces between glued words in common patterns
    """
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")

    # Joined words due to line wraps: "сло-\nво" -> "слово"
    normalized = re.sub(r"(?<=\w)-\n(?=\w)", "", normalized)

    # Keep empty line as paragraph delimiter, but join single line wraps with spaces.
    normalized = re.sub(r"(?<!\n)\n(?!\n)", " ", normalized)

    # Insert space in common glued boundaries.
    normalized = re.sub(r"(?<=[a-zа-яё])(?=[A-ZА-ЯЁ])", " ", normalized)
    normalized = re.sub(r"(?<=[A-Za-zА-Яа-яЁё])(?=\d)", " ", normalized)
    normalized = re.sub(r"(?<=\d)(?=[A-Za-zА-Яа-яЁё])", " ", normalized)

    # Normalize whitespace.
    normalized = re.sub(r"[ \t]+", " ", normalized)
    normalized = re.sub(r" *\n{2,} *", "\n\n", normalized)

    return normalized.strip()


def _run_ocrmypdf(
    input_pdf_path: Path,
    output_pdf_path: Path,
    sidecar_path: Optional[Path] = None,
) -> None:
    command = _resolve_ocrmypdf_command()
    if not command:
        raise RuntimeError(
            "Обнаружен сканированный PDF, но OCR не настроен. "
            "Установите ocrmypdf и Tesseract, затем повторите попытку."
        )

    extra_path_dirs = _collect_windows_ocr_tool_dirs()
    env = _build_ocr_env(extra_path_dirs)
    ocr_lang = _resolve_ocr_lang(extra_path_dirs)

    cmd = [
        *command,
        "--force-ocr",
        "--rotate-pages",
        "--deskew",
        "--pdf-renderer",
        "auto",
        "--optimize",
        "0",
        "--tesseract-pagesegmode",
        "6",
        "-l",
        ocr_lang,
        str(input_pdf_path),
        str(output_pdf_path),
    ]
    if sidecar_path is not None:
        cmd.extend(["--sidecar", str(sidecar_path)])
    process = subprocess.run(cmd, capture_output=True, text=True, env=env)
    if process.returncode != 0:
        details = (process.stderr or process.stdout or "неизвестная ошибка OCR").strip()
        raise RuntimeError(f"Ошибка OCR: {details}")


def _build_docx_from_sidecar_text(sidecar_path: Path, docx_path: Path) -> int:
    """
    Build DOCX from OCR sidecar text. This is the most reliable OCR output source.
    """
    if not sidecar_path.exists():
        return 0

    raw_text = sidecar_path.read_text(encoding="utf-8", errors="ignore").strip()
    if not raw_text:
        return 0

    from docx import Document

    document = Document()
    total_chars = 0

    pages = [page.strip() for page in raw_text.split("\f")]
    for page in pages:
        normalized_page = _normalize_ocr_text(page)
        if not normalized_page:
            continue

        if document.paragraphs:
            document.add_page_break()

        for paragraph in (chunk.strip() for chunk in normalized_page.split("\n\n")):
            if paragraph:
                document.add_paragraph(paragraph)
                total_chars += len(paragraph)

    if total_chars == 0:
        return 0

    document.save(str(docx_path))
    return total_chars


def _build_docx_from_pdf_text(pdf_path: Path, docx_path: Path) -> int:
    """
    Build text-first DOCX from PDF text layer.
    Used for scanned PDFs after OCR, because pdf2docx may still keep image-only pages.
    """
    _ensure_pymupdf_compat()
    import fitz  # PyMuPDF
    from docx import Document

    document = Document()
    total_chars = 0

    with fitz.open(str(pdf_path)) as pdf_doc:
        for page_index in range(pdf_doc.page_count):
            page_text = pdf_doc.load_page(page_index).get_text("text").strip()
            normalized_page = _normalize_ocr_text(page_text)
            if not normalized_page:
                continue

            total_chars += len(normalized_page)
            if document.paragraphs:
                document.add_page_break()

            for paragraph in (chunk.strip() for chunk in normalized_page.split("\n\n")):
                if paragraph:
                    document.add_paragraph(paragraph)

    if total_chars == 0:
        return 0

    document.save(str(docx_path))
    return total_chars


def _convert_with_pipeline(pdf_path: Path, docx_path: Path, task_uuid: str) -> None:
    pdf_kind = _detect_pdf_kind(pdf_path)

    if pdf_kind == TEXT:
        logger.info("Task %s: detected text PDF, using direct pdf2docx", task_uuid)
        _do_convert(str(pdf_path), str(docx_path))
        return

    if not settings.ENABLE_SCANNED_OCR:
        raise RuntimeError(
            "Обнаружен сканированный PDF, но OCR отключён "
            "(установите ENABLE_SCANNED_OCR=true)."
        )

    logger.info("Task %s: detected scanned PDF, running OCR pipeline", task_uuid)
    ocr_pdf_path = pdf_path.with_name(f"{pdf_path.stem}.ocr.pdf")
    ocr_sidecar_path = pdf_path.with_name(f"{pdf_path.stem}.ocr.txt")
    try:
        _run_ocrmypdf(pdf_path, ocr_pdf_path, sidecar_path=ocr_sidecar_path)

        sidecar_chars = _build_docx_from_sidecar_text(ocr_sidecar_path, docx_path)
        if sidecar_chars > 0:
            logger.info(
                "Task %s: OCR sidecar text extracted (%s chars)",
                task_uuid,
                sidecar_chars,
            )
            return

        pdf_chars = _build_docx_from_pdf_text(ocr_pdf_path, docx_path)
        if pdf_chars > 0:
            logger.info(
                "Task %s: OCR PDF text layer extracted (%s chars)",
                task_uuid,
                pdf_chars,
            )
            return

        raise RuntimeError(
            "OCR завершён, но текст не извлечён. "
            "Проверьте качество скана (размытие/DPI), языковые пакеты и OCR_LANG."
        )
    finally:
        try:
            os.remove(ocr_pdf_path)
        except OSError:
            pass
        try:
            os.remove(ocr_sidecar_path)
        except OSError:
            pass


def _convert_pdf_to_jpg_archive(pdf_path: Path, zip_path: Path, task_uuid: str) -> None:
    _ensure_pymupdf_compat()
    import fitz

    logger.info("Task %s: converting PDF to JPG archive", task_uuid)

    with fitz.open(str(pdf_path)) as pdf_doc:
        if pdf_doc.page_count == 0:
            raise RuntimeError("PDF не содержит страниц для конвертации")

        with ZipFile(zip_path, "w", compression=ZIP_DEFLATED) as archive:
            for page_index in range(pdf_doc.page_count):
                page = pdf_doc.load_page(page_index)
                pix = page.get_pixmap(matrix=fitz.Matrix(2.0, 2.0), alpha=False)
                jpg_name = f"page_{page_index + 1:03d}.jpg"
                jpg_bytes = pix.tobytes("jpg")
                archive.writestr(jpg_name, jpg_bytes)


def _convert_jpg_to_pdf(jpg_path: Path, pdf_path: Path, task_uuid: str) -> None:
    _ensure_pymupdf_compat()
    import fitz

    logger.info("Task %s: converting JPG to PDF", task_uuid)

    image_doc = fitz.open(str(jpg_path))
    try:
        pdf_bytes = image_doc.convert_to_pdf()
    finally:
        image_doc.close()

    pdf_doc = fitz.open("pdf", pdf_bytes)
    try:
        pdf_doc.save(str(pdf_path))
    finally:
        pdf_doc.close()


def _convert_jpg_list_to_pdf(jpg_paths: list[Path], pdf_path: Path, task_uuid: str) -> None:
    _ensure_pymupdf_compat()
    import fitz

    logger.info("Task %s: converting %s JPG files to one PDF", task_uuid, len(jpg_paths))

    if not jpg_paths:
        raise RuntimeError("Не передано ни одного JPG-файла для конвертации")

    merged_pdf = fitz.open()
    try:
        for jpg_path in jpg_paths:
            image_doc = fitz.open(str(jpg_path))
            try:
                pdf_bytes = image_doc.convert_to_pdf()
            finally:
                image_doc.close()

            image_pdf = fitz.open("pdf", pdf_bytes)
            try:
                merged_pdf.insert_pdf(image_pdf)
            finally:
                image_pdf.close()

        if merged_pdf.page_count == 0:
            raise RuntimeError("Не удалось собрать PDF из изображений")

        merged_pdf.save(str(pdf_path))
    finally:
        merged_pdf.close()


def _convert_heic_to_jpg(heic_path: Path, jpg_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting HEIC to JPG", task_uuid)

    try:
        import pillow_heif
        from PIL import Image

        pillow_heif.register_heif_opener()
        with Image.open(str(heic_path)) as image:
            rgb = image.convert("RGB")
            rgb.save(str(jpg_path), format="JPEG", quality=92, optimize=True)
        return
    except Exception as exc:
        logger.warning("Task %s: PIL/pillow-heif HEIC conversion failed: %s", task_uuid, exc)

    _run_ffmpeg(
        [
            "-y",
            "-i",
            str(heic_path),
            "-frames:v",
            "1",
            "-q:v",
            "2",
            str(jpg_path),
        ],
        "Ошибка FFmpeg (HEIC -> JPG)",
    )


def _convert_png_to_webp(png_path: Path, webp_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting PNG to WEBP", task_uuid)
    from PIL import Image

    with Image.open(str(png_path)) as image:
        converted = image.convert("RGBA" if image.mode in {"RGBA", "LA", "P"} else "RGB")
        converted.save(str(webp_path), format="WEBP", quality=90, method=6)


def _convert_svg_to_png(svg_path: Path, png_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting SVG to PNG", task_uuid)

    try:
        from cairosvg import svg2png
    except ImportError as exc:
        raise RuntimeError(
            "Для SVG -> PNG требуется библиотека cairosvg (pip install cairosvg)."
        ) from exc

    svg2png(url=str(svg_path), write_to=str(png_path))


def _convert_png_to_ico(png_path: Path, ico_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting PNG to ICO", task_uuid)
    from PIL import Image

    with Image.open(str(png_path)) as image:
        icon = image.convert("RGBA")
        icon.save(
            str(ico_path),
            format="ICO",
            sizes=[(16, 16), (32, 32), (48, 48), (64, 64), (128, 128), (256, 256)],
        )


def _convert_psd_ai_to_png(source_path: Path, png_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting PSD/AI to PNG", task_uuid)

    try:
        from PIL import Image

        with Image.open(str(source_path)) as image:
            converted = image.convert("RGBA")
            converted.save(str(png_path), format="PNG")
        return
    except Exception as exc:
        logger.warning("Task %s: PIL PSD/AI -> PNG conversion failed: %s", task_uuid, exc)

    _run_imagemagick(
        [
            str(source_path),
            str(png_path),
        ],
        "Ошибка ImageMagick (PSD/AI -> PNG)",
    )


def _convert_psd_ai_to_jpg(source_path: Path, jpg_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting PSD/AI to JPG", task_uuid)

    try:
        from PIL import Image

        with Image.open(str(source_path)) as image:
            rgb = image.convert("RGB")
            rgb.save(str(jpg_path), format="JPEG", quality=92, optimize=True)
        return
    except Exception as exc:
        logger.warning("Task %s: PIL PSD/AI -> JPG conversion failed: %s", task_uuid, exc)

    _run_imagemagick(
        [
            str(source_path),
            "-background",
            "white",
            "-alpha",
            "remove",
            str(jpg_path),
        ],
        "Ошибка ImageMagick (PSD/AI -> JPG)",
    )


def _convert_video_to_mp4(video_path: Path, mp4_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting video to MP4", task_uuid)
    _run_ffmpeg(
        [
            "-y",
            "-i",
            str(video_path),
            "-c:v",
            "libx264",
            "-preset",
            "medium",
            "-crf",
            "23",
            "-c:a",
            "aac",
            "-b:a",
            "128k",
            "-movflags",
            "+faststart",
            str(mp4_path),
        ],
        "Ошибка FFmpeg (Видео -> MP4)",
    )


def _convert_gif_to_mp4(gif_path: Path, mp4_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting GIF to MP4", task_uuid)
    _run_ffmpeg(
        [
            "-y",
            "-i",
            str(gif_path),
            "-vf",
            "fps=24,scale=trunc(iw/2)*2:trunc(ih/2)*2:flags=lanczos",
            "-c:v",
            "libx264",
            "-pix_fmt",
            "yuv420p",
            "-movflags",
            "+faststart",
            "-an",
            str(mp4_path),
        ],
        "Ошибка FFmpeg (GIF -> MP4)",
    )


def _convert_video_to_mp3(video_path: Path, mp3_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: extracting MP3 from video", task_uuid)
    _run_ffmpeg(
        [
            "-y",
            "-i",
            str(video_path),
            "-vn",
            "-acodec",
            "libmp3lame",
            "-q:a",
            "2",
            str(mp3_path),
        ],
        "Ошибка FFmpeg (Видео -> MP3)",
    )


def _convert_video_to_avi(video_path: Path, avi_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting video to AVI", task_uuid)
    _run_ffmpeg(
        [
            "-y",
            "-i",
            str(video_path),
            "-c:v",
            "mpeg4",
            "-q:v",
            "4",
            "-c:a",
            "mp3",
            str(avi_path),
        ],
        "Ошибка FFmpeg (Видео -> AVI)",
    )


def _convert_video_to_mov(video_path: Path, mov_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting video to MOV", task_uuid)
    _run_ffmpeg(
        [
            "-y",
            "-i",
            str(video_path),
            "-c:v",
            "libx264",
            "-c:a",
            "aac",
            "-movflags",
            "+faststart",
            str(mov_path),
        ],
        "Ошибка FFmpeg (Видео -> MOV)",
    )


def _pack_files_to_zip(
    file_items: list[tuple[str, Path]],
    output_zip_path: Path,
    task_uuid: str,
) -> None:
    logger.info("Task %s: packing %s files to ZIP", task_uuid, len(file_items))
    if not file_items:
        raise RuntimeError("Не переданы файлы для архивации")
    _zip_files(file_items, output_zip_path)


def _pack_files_to_7z(
    file_items: list[tuple[str, Path]],
    output_7z_path: Path,
    task_uuid: str,
) -> None:
    logger.info("Task %s: packing %s files to 7Z", task_uuid, len(file_items))
    if not file_items:
        raise RuntimeError("Не переданы файлы для архивации")

    try:
        import py7zr
    except ImportError as exc:
        raise RuntimeError("Для создания 7Z требуется библиотека py7zr.") from exc

    used: set[str] = set()
    with py7zr.SevenZipFile(output_7z_path, "w") as archive:
        for original_name, file_path in file_items:
            arcname = _unique_name(original_name, used)
            archive.write(file_path, arcname=arcname)


def _pack_files_to_rar(
    file_items: list[tuple[str, Path]],
    output_rar_path: Path,
    task_uuid: str,
) -> None:
    logger.info("Task %s: packing %s files to RAR", task_uuid, len(file_items))
    if not file_items:
        raise RuntimeError("Не переданы файлы для архивации")

    rar_cmd = _resolve_rar_command()
    if not rar_cmd:
        raise RuntimeError("RAR-архиватор не найден. Установите WinRAR/rar CLI.")

    with tempfile.TemporaryDirectory(prefix="fmtswap_rar_pack_") as tmp_dir_name:
        tmp_dir = Path(tmp_dir_name)
        used: set[str] = set()
        staged_paths: list[Path] = []
        for original_name, file_path in file_items:
            safe_name = _unique_name(original_name, used)
            staged_path = tmp_dir / safe_name
            staged_path.parent.mkdir(parents=True, exist_ok=True)
            shutil.copy2(file_path, staged_path)
            staged_paths.append(staged_path)

        cmd = [
            rar_cmd,
            "a",
            "-ep1",
            str(output_rar_path),
            *[str(path) for path in staged_paths],
        ]
        process = subprocess.run(cmd, capture_output=True, text=True)
        if process.returncode != 0:
            details = (process.stderr or process.stdout or "ошибка упаковки RAR").strip()
            raise RuntimeError(f"Ошибка RAR (упаковка): {details}")


def _unpack_zip_to_zip(input_zip_path: Path, output_zip_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: unpacking ZIP", task_uuid)
    with tempfile.TemporaryDirectory(prefix="fmtswap_zip_unpack_") as tmp_dir_name:
        tmp_dir = Path(tmp_dir_name)
        with ZipFile(input_zip_path, "r") as archive:
            for member in archive.infolist():
                if member.is_dir():
                    continue
                safe_member = _sanitize_member_name(member.filename)
                if safe_member is None:
                    continue
                target = (tmp_dir / safe_member).resolve()
                if not str(target).startswith(str(tmp_dir.resolve())):
                    continue
                target.parent.mkdir(parents=True, exist_ok=True)
                with archive.open(member, "r") as src, open(target, "wb") as dst:
                    shutil.copyfileobj(src, dst)

        _zip_directory(tmp_dir, output_zip_path)


def _unpack_7z_to_zip(input_7z_path: Path, output_zip_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: unpacking 7Z", task_uuid)
    try:
        import py7zr
    except ImportError as exc:
        raise RuntimeError("Для распаковки 7Z требуется библиотека py7zr.") from exc

    with tempfile.TemporaryDirectory(prefix="fmtswap_7z_unpack_") as tmp_dir_name:
        tmp_dir = Path(tmp_dir_name)
        with py7zr.SevenZipFile(input_7z_path, "r") as archive:
            archive.extractall(path=tmp_dir)
        _zip_directory(tmp_dir, output_zip_path)


def _unpack_rar_to_zip(input_rar_path: Path, output_zip_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: unpacking RAR", task_uuid)
    try:
        import rarfile
    except ImportError as exc:
        raise RuntimeError("Для распаковки RAR требуется библиотека rarfile.") from exc

    with tempfile.TemporaryDirectory(prefix="fmtswap_rar_unpack_") as tmp_dir_name:
        tmp_dir = Path(tmp_dir_name)
        with rarfile.RarFile(input_rar_path) as archive:
            archive.extractall(path=tmp_dir)
        _zip_directory(tmp_dir, output_zip_path)


def _prepare_image_for_compression(input_path: Path):
    from PIL import Image
    try:
        import pillow_heif

        pillow_heif.register_heif_opener()
    except Exception:
        pass

    image = Image.open(str(input_path))

    # Resize very large images for practical web/email size.
    max_width = 1920
    if image.width > max_width:
        ratio = max_width / float(image.width)
        image = image.resize((max_width, max(1, int(image.height * ratio))), Image.Resampling.LANCZOS)
    return image


def _compress_image_to_jpg(input_path: Path, output_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: compressing image to JPG", task_uuid)
    from PIL import Image

    image = _prepare_image_for_compression(input_path)
    try:
        if image.mode in {"RGBA", "LA", "P"}:
            rgba = image.convert("RGBA")
            canvas = rgba.getchannel("A")
            flattened = Image.new("RGB", rgba.size, "white")
            flattened.paste(rgba, mask=canvas)
            image = flattened
        else:
            image = image.convert("RGB")
        image.save(str(output_path), format="JPEG", quality=78, optimize=True, progressive=True)
    finally:
        image.close()


def _compress_image_to_webp(input_path: Path, output_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: compressing image to WEBP", task_uuid)
    image = _prepare_image_for_compression(input_path)
    try:
        if image.mode not in {"RGB", "RGBA"}:
            image = image.convert("RGB")
        image.save(str(output_path), format="WEBP", quality=80, method=6)
    finally:
        image.close()


def _compress_image_to_avif(input_path: Path, output_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: compressing image to AVIF", task_uuid)
    try:
        image = _prepare_image_for_compression(input_path)
        try:
            image.save(str(output_path), format="AVIF", quality=50)
            return
        finally:
            image.close()
    except Exception as exc:
        logger.warning("Task %s: AVIF via PIL failed, fallback to FFmpeg: %s", task_uuid, exc)

    _run_ffmpeg(
        [
            "-y",
            "-i",
            str(input_path),
            "-c:v",
            "libaom-av1",
            "-still-picture",
            "1",
            "-crf",
            "35",
            "-b:v",
            "0",
            str(output_path),
        ],
        "Ошибка FFmpeg (Изображение -> AVIF)",
    )


def _compress_pdf(input_pdf_path: Path, output_pdf_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: compressing PDF", task_uuid)
    _run_ghostscript(
        [
            "-sDEVICE=pdfwrite",
            "-dCompatibilityLevel=1.5",
            "-dPDFSETTINGS=/ebook",
            "-dNOPAUSE",
            "-dQUIET",
            "-dBATCH",
            "-dDetectDuplicateImages=true",
            "-dDownsampleColorImages=true",
            "-dColorImageResolution=150",
            f"-sOutputFile={output_pdf_path}",
            str(input_pdf_path),
        ],
        "Ошибка Ghostscript (PDF сжатие)",
    )


def _compress_office_to_zip(input_path: Path, output_zip_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: compressing office document to ZIP", task_uuid)
    with ZipFile(output_zip_path, "w", compression=ZIP_DEFLATED, compresslevel=9) as archive:
        archive.write(input_path, arcname=input_path.name)


def _compress_video_to_mp4(input_path: Path, output_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: compressing video to MP4", task_uuid)
    _run_ffmpeg(
        [
            "-y",
            "-i",
            str(input_path),
            "-vf",
            "scale='min(1280,iw)':-2",
            "-c:v",
            "libx264",
            "-preset",
            "medium",
            "-crf",
            "28",
            "-maxrate",
            "1800k",
            "-bufsize",
            "3600k",
            "-c:a",
            "aac",
            "-b:a",
            "96k",
            "-movflags",
            "+faststart",
            str(output_path),
        ],
        "Ошибка FFmpeg (Видео сжатие)",
    )


def _compress_audio_to_mp3(input_path: Path, output_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: compressing audio to MP3", task_uuid)
    _run_ffmpeg(
        [
            "-y",
            "-i",
            str(input_path),
            "-vn",
            "-ac",
            "2",
            "-ar",
            "44100",
            "-b:a",
            "128k",
            str(output_path),
        ],
        "Ошибка FFmpeg (Аудио сжатие)",
    )


def _resolve_unicode_font_file() -> Optional[str]:
    """
    Try to find a system font that supports Cyrillic/Unicode text for fallback rendering.
    """
    candidates: list[Path] = []

    if os.name == "nt":
        win_fonts = Path(r"C:\Windows\Fonts")
        candidates.extend(
            [
                win_fonts / "arial.ttf",
                win_fonts / "calibri.ttf",
                win_fonts / "times.ttf",
                win_fonts / "tahoma.ttf",
                win_fonts / "segoeui.ttf",
            ]
        )
    else:
        candidates.extend(
            [
                Path("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
                Path("/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf"),
                Path("/usr/share/fonts/truetype/freefont/FreeSans.ttf"),
            ]
        )

    for font_path in candidates:
        if font_path.exists():
            return str(font_path)

    return None


def _resolve_soffice_command() -> Optional[str]:
    soffice = shutil.which("soffice")
    if soffice:
        return soffice

    if os.name == "nt":
        candidates = [
            Path(r"C:\Program Files\LibreOffice\program\soffice.exe"),
            Path(r"C:\Program Files (x86)\LibreOffice\program\soffice.exe"),
        ]
        for candidate in candidates:
            if candidate.exists():
                return str(candidate)

    return None


def _convert_word_with_libreoffice(word_path: Path, pdf_path: Path) -> None:
    soffice_cmd = _resolve_soffice_command()
    if not soffice_cmd:
        raise RuntimeError("LibreOffice (soffice) не найден в системе")

    with tempfile.TemporaryDirectory(prefix="fmtswap_soffice_") as tmp_dir:
        cmd = [
            soffice_cmd,
            "--headless",
            "--nologo",
            "--nolockcheck",
            "--nodefault",
            "--nofirststartwizard",
            "--convert-to",
            "pdf:writer_pdf_Export",
            "--outdir",
            tmp_dir,
            str(word_path),
        ]
        process = subprocess.run(cmd, capture_output=True, text=True, timeout=240)
        if process.returncode != 0:
            details = (process.stderr or process.stdout or "ошибка LibreOffice").strip()
            raise RuntimeError(details)

        converted_path = Path(tmp_dir) / f"{word_path.stem}.pdf"
        if not converted_path.exists():
            generated = sorted(Path(tmp_dir).glob("*.pdf"))
            if not generated:
                raise RuntimeError("LibreOffice не сгенерировал PDF-файл")
            converted_path = generated[0]

        if pdf_path.exists():
            pdf_path.unlink()
        shutil.move(str(converted_path), str(pdf_path))


def _convert_word_with_docx2pdf(word_path: Path, pdf_path: Path) -> None:
    if importlib.util.find_spec("docx2pdf") is None:
        raise RuntimeError("Модуль docx2pdf не установлен")

    from docx2pdf import convert as docx2pdf_convert

    docx2pdf_convert(str(word_path), str(pdf_path))
    if not pdf_path.exists():
        raise RuntimeError("docx2pdf не сгенерировал PDF-файл")


def _convert_docx_to_pdf_fallback(docx_path: Path, pdf_path: Path, task_uuid: str) -> None:
    _ensure_pymupdf_compat()
    import fitz
    from docx import Document

    logger.info("Task %s: converting WORD to PDF via text fallback", task_uuid)

    doc = Document(str(docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs]

    if not any(paragraphs):
        raise RuntimeError("Документ Word не содержит текста для конвертации")

    pdf = fitz.open()
    margin = 54
    font_size = 11
    line_height = 16
    max_chars = 95
    font_file = _resolve_unicode_font_file()

    if font_file:
        logger.info("Task %s: WORD fallback uses system font %s", task_uuid, font_file)

    page = pdf.new_page()
    y = margin

    def ensure_space(lines_count: int) -> None:
        nonlocal page, y
        required_height = lines_count * line_height
        if y + required_height > page.rect.height - margin:
            page = pdf.new_page()
            y = margin

    for paragraph in paragraphs:
        wrapped_lines = textwrap.wrap(paragraph, width=max_chars) if paragraph else [""]
        ensure_space(len(wrapped_lines) + 1)

        for line in wrapped_lines:
            insert_kwargs = {
                "fontsize": font_size,
                "color": (0, 0, 0),
            }
            if font_file:
                insert_kwargs.update({"fontname": "fmtswap_unicode", "fontfile": font_file})
            else:
                # Fallback when no Unicode TTF found.
                insert_kwargs.update({"fontname": "helv", "encoding": fitz.TEXT_ENCODING_CYRILLIC})

            page.insert_text((margin, y), line, **insert_kwargs)
            y += line_height

        y += line_height // 2

    pdf.save(str(pdf_path))
    pdf.close()


def _convert_word_to_pdf(word_path: Path, pdf_path: Path, task_uuid: str) -> None:
    logger.info("Task %s: converting WORD to PDF", task_uuid)
    ext = word_path.suffix.lower()
    errors: list[str] = []

    try:
        _convert_word_with_libreoffice(word_path, pdf_path)
        return
    except Exception as exc:
        errors.append(f"LibreOffice: {exc}")

    if ext in {".docx", ".docm", ".dotx", ".dotm"}:
        try:
            _convert_word_with_docx2pdf(word_path, pdf_path)
            return
        except Exception as exc:
            errors.append(f"docx2pdf: {exc}")

        try:
            _convert_docx_to_pdf_fallback(word_path, pdf_path, task_uuid)
            return
        except Exception as exc:
            errors.append(f"fallback: {exc}")

    raise RuntimeError(
        "Не удалось конвертировать Word в PDF. "
        "Рекомендуется установить LibreOffice для качественной конвертации. "
        f"Подробности: {' | '.join(errors)}"
    )


def get_supported_conversion_types() -> set[str]:
    return {
        PDF_TO_DOCX,
        PDF_TO_JPG,
        JPG_TO_PDF,
        WORD_TO_PDF,
        HEIC_TO_JPG,
        PNG_TO_WEBP,
        SVG_TO_PNG,
        PNG_TO_ICO,
        PSD_AI_TO_PNG,
        PSD_AI_TO_JPG,
        VIDEO_TO_MP4,
        GIF_TO_MP4,
        VIDEO_TO_MP3,
        VIDEO_TO_AVI,
        VIDEO_TO_MOV,
        ARCH_ZIP_PACK,
        ARCH_7Z_PACK,
        ARCH_RAR_PACK,
        ARCH_ZIP_UNPACK,
        ARCH_7Z_UNPACK,
        ARCH_RAR_UNPACK,
        COMP_IMG_JPG,
        COMP_IMG_WEBP,
        COMP_IMG_AVIF,
        COMP_PDF,
        COMP_OFFICE_ZIP,
        COMP_VIDEO_MP4,
        COMP_AUDIO_MP3,
    }


def get_input_extensions(conversion_type: str) -> tuple[str, ...]:
    mapping = {
        PDF_TO_DOCX: (".pdf",),
        PDF_TO_JPG: (".pdf",),
        JPG_TO_PDF: (".jpg", ".jpeg", ".jfif"),
        WORD_TO_PDF: (".docx", ".doc", ".docm"),
        HEIC_TO_JPG: (".heic", ".heif"),
        PNG_TO_WEBP: (".png",),
        SVG_TO_PNG: (".svg",),
        PNG_TO_ICO: (".png",),
        PSD_AI_TO_PNG: (".psd", ".ai"),
        PSD_AI_TO_JPG: (".psd", ".ai"),
        VIDEO_TO_MP4: VIDEO_INPUT_EXTENSIONS,
        GIF_TO_MP4: (".gif",),
        VIDEO_TO_MP3: VIDEO_INPUT_EXTENSIONS,
        VIDEO_TO_AVI: VIDEO_INPUT_EXTENSIONS,
        VIDEO_TO_MOV: VIDEO_INPUT_EXTENSIONS,
        ARCH_ZIP_PACK: ("*",),
        ARCH_7Z_PACK: ("*",),
        ARCH_RAR_PACK: ("*",),
        ARCH_ZIP_UNPACK: (".zip",),
        ARCH_7Z_UNPACK: (".7z",),
        ARCH_RAR_UNPACK: (".rar",),
        COMP_IMG_JPG: (".jpg", ".jpeg", ".png", ".heic", ".heif"),
        COMP_IMG_WEBP: (".jpg", ".jpeg", ".png", ".heic", ".heif"),
        COMP_IMG_AVIF: (".jpg", ".jpeg", ".png", ".heic", ".heif"),
        COMP_PDF: (".pdf",),
        COMP_OFFICE_ZIP: (".doc", ".docx", ".xls", ".xlsx", ".ppt", ".pptx"),
        COMP_VIDEO_MP4: VIDEO_INPUT_EXTENSIONS,
        COMP_AUDIO_MP3: (*VIDEO_INPUT_EXTENSIONS, *AUDIO_INPUT_EXTENSIONS),
    }
    return mapping.get(conversion_type, tuple())


def get_output_extension(conversion_type: str) -> str:
    mapping = {
        PDF_TO_DOCX: ".docx",
        PDF_TO_JPG: ".zip",
        JPG_TO_PDF: ".pdf",
        WORD_TO_PDF: ".pdf",
        HEIC_TO_JPG: ".jpg",
        PNG_TO_WEBP: ".webp",
        SVG_TO_PNG: ".png",
        PNG_TO_ICO: ".ico",
        PSD_AI_TO_PNG: ".png",
        PSD_AI_TO_JPG: ".jpg",
        VIDEO_TO_MP4: ".mp4",
        GIF_TO_MP4: ".mp4",
        VIDEO_TO_MP3: ".mp3",
        VIDEO_TO_AVI: ".avi",
        VIDEO_TO_MOV: ".mov",
        ARCH_ZIP_PACK: ".zip",
        ARCH_7Z_PACK: ".7z",
        ARCH_RAR_PACK: ".rar",
        ARCH_ZIP_UNPACK: ".zip",
        ARCH_7Z_UNPACK: ".zip",
        ARCH_RAR_UNPACK: ".zip",
        COMP_IMG_JPG: ".jpg",
        COMP_IMG_WEBP: ".webp",
        COMP_IMG_AVIF: ".avif",
        COMP_PDF: ".pdf",
        COMP_OFFICE_ZIP: ".zip",
        COMP_VIDEO_MP4: ".mp4",
        COMP_AUDIO_MP3: ".mp3",
    }
    return mapping.get(conversion_type, ".bin")


def get_output_media_type(conversion_type: str) -> str:
    mapping = {
        PDF_TO_DOCX: "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        PDF_TO_JPG: "application/zip",
        JPG_TO_PDF: "application/pdf",
        WORD_TO_PDF: "application/pdf",
        HEIC_TO_JPG: "image/jpeg",
        PNG_TO_WEBP: "image/webp",
        SVG_TO_PNG: "image/png",
        PNG_TO_ICO: "image/x-icon",
        PSD_AI_TO_PNG: "image/png",
        PSD_AI_TO_JPG: "image/jpeg",
        VIDEO_TO_MP4: "video/mp4",
        GIF_TO_MP4: "video/mp4",
        VIDEO_TO_MP3: "audio/mpeg",
        VIDEO_TO_AVI: "video/x-msvideo",
        VIDEO_TO_MOV: "video/quicktime",
        ARCH_ZIP_PACK: "application/zip",
        ARCH_7Z_PACK: "application/x-7z-compressed",
        ARCH_RAR_PACK: "application/vnd.rar",
        ARCH_ZIP_UNPACK: "application/zip",
        ARCH_7Z_UNPACK: "application/zip",
        ARCH_RAR_UNPACK: "application/zip",
        COMP_IMG_JPG: "image/jpeg",
        COMP_IMG_WEBP: "image/webp",
        COMP_IMG_AVIF: "image/avif",
        COMP_PDF: "application/pdf",
        COMP_OFFICE_ZIP: "application/zip",
        COMP_VIDEO_MP4: "video/mp4",
        COMP_AUDIO_MP3: "audio/mpeg",
    }
    return mapping.get(conversion_type, "application/octet-stream")


def _resolve_runner(conversion_type: str):
    mapping = {
        PDF_TO_DOCX: _convert_with_pipeline,
        PDF_TO_JPG: _convert_pdf_to_jpg_archive,
        JPG_TO_PDF: _convert_jpg_to_pdf,
        WORD_TO_PDF: _convert_word_to_pdf,
        HEIC_TO_JPG: _convert_heic_to_jpg,
        PNG_TO_WEBP: _convert_png_to_webp,
        SVG_TO_PNG: _convert_svg_to_png,
        PNG_TO_ICO: _convert_png_to_ico,
        PSD_AI_TO_PNG: _convert_psd_ai_to_png,
        PSD_AI_TO_JPG: _convert_psd_ai_to_jpg,
        VIDEO_TO_MP4: _convert_video_to_mp4,
        GIF_TO_MP4: _convert_gif_to_mp4,
        VIDEO_TO_MP3: _convert_video_to_mp3,
        VIDEO_TO_AVI: _convert_video_to_avi,
        VIDEO_TO_MOV: _convert_video_to_mov,
        ARCH_ZIP_UNPACK: _unpack_zip_to_zip,
        ARCH_7Z_UNPACK: _unpack_7z_to_zip,
        ARCH_RAR_UNPACK: _unpack_rar_to_zip,
        COMP_IMG_JPG: _compress_image_to_jpg,
        COMP_IMG_WEBP: _compress_image_to_webp,
        COMP_IMG_AVIF: _compress_image_to_avif,
        COMP_PDF: _compress_pdf,
        COMP_OFFICE_ZIP: _compress_office_to_zip,
        COMP_VIDEO_MP4: _compress_video_to_mp4,
        COMP_AUDIO_MP3: _compress_audio_to_mp3,
    }
    return mapping.get(conversion_type)


def _resolve_batch_runner(conversion_type: str):
    mapping = {
        ARCH_ZIP_PACK: _pack_files_to_zip,
        ARCH_7Z_PACK: _pack_files_to_7z,
        ARCH_RAR_PACK: _pack_files_to_rar,
    }
    return mapping.get(conversion_type)


async def convert_file(
    task_uuid: str,
    source_bytes: bytes,
    original_filename: str,
    conversion_type: str,
    db: Session,
) -> None:
    """
    Save uploaded file, run conversion asynchronously, and update DB status.
    """
    upload_dir, output_dir = _ensure_dirs()

    output_ext = get_output_extension(conversion_type)
    input_ext = Path(original_filename).suffix.lower() or ".bin"

    input_filename = f"{task_uuid}{input_ext}"
    output_filename = f"{task_uuid}{output_ext}"
    input_path = upload_dir / input_filename
    output_path = output_dir / output_filename

    with open(input_path, "wb") as f:
        f.write(source_bytes)

    import uuid as _uuid

    try:
        uid = _uuid.UUID(task_uuid)
    except (ValueError, AttributeError):
        uid = task_uuid

    task = db.query(ConversionTask).filter(ConversionTask.task_uuid == uid).first()
    if not task:
        logger.error("Task %s not found in DB", task_uuid)
        return

    task.status = TaskStatus.PROCESSING
    task.updated_at = datetime.now(timezone.utc)
    db.commit()

    try:
        runner = _resolve_runner(conversion_type)
        if runner is None:
            raise RuntimeError(f"Неподдерживаемый тип конвертации: {conversion_type}")

        loop = asyncio.get_running_loop()
        await loop.run_in_executor(None, runner, input_path, output_path, task_uuid)

        task.status = TaskStatus.DONE
        task.output_filename = output_filename
        task.updated_at = datetime.now(timezone.utc)
        db.commit()
        logger.info("Task %s completed successfully", task_uuid)

    except Exception as exc:
        logger.exception("Conversion failed for task %s: %s", task_uuid, exc)
        task.status = TaskStatus.FAILED
        task.error_message = str(exc)
        task.updated_at = datetime.now(timezone.utc)
        db.commit()

    finally:
        try:
            os.remove(input_path)
        except OSError:
            pass


async def convert_images_to_pdf(
    task_uuid: str,
    images: list[tuple[str, bytes]],
    db: Session,
) -> None:
    """
    Convert multiple JPG/JPEG images into a single multi-page PDF.
    """
    upload_dir, output_dir = _ensure_dirs()

    output_filename = f"{task_uuid}.pdf"
    output_path = output_dir / output_filename

    image_paths: list[Path] = []
    for idx, (filename, payload) in enumerate(images, start=1):
        ext = Path(filename).suffix.lower() or ".jpg"
        image_path = upload_dir / f"{task_uuid}_{idx:03d}{ext}"
        with open(image_path, "wb") as f:
            f.write(payload)
        image_paths.append(image_path)

    import uuid as _uuid

    try:
        uid = _uuid.UUID(task_uuid)
    except (ValueError, AttributeError):
        uid = task_uuid

    task = db.query(ConversionTask).filter(ConversionTask.task_uuid == uid).first()
    if not task:
        logger.error("Task %s not found in DB", task_uuid)
        return

    task.status = TaskStatus.PROCESSING
    task.updated_at = datetime.now(timezone.utc)
    db.commit()

    try:
        loop = asyncio.get_running_loop()
        await loop.run_in_executor(None, _convert_jpg_list_to_pdf, image_paths, output_path, task_uuid)

        task.status = TaskStatus.DONE
        task.output_filename = output_filename
        task.updated_at = datetime.now(timezone.utc)
        db.commit()
        logger.info("Task %s completed successfully", task_uuid)

    except Exception as exc:
        logger.exception("Multi-image conversion failed for task %s: %s", task_uuid, exc)
        task.status = TaskStatus.FAILED
        task.error_message = str(exc)
        task.updated_at = datetime.now(timezone.utc)
        db.commit()

    finally:
        for image_path in image_paths:
            try:
                os.remove(image_path)
            except OSError:
                pass


async def convert_files_batch(
    task_uuid: str,
    files: list[tuple[str, bytes]],
    conversion_type: str,
    db: Session,
) -> None:
    """
    Process multi-file conversions (archive packers).
    """
    upload_dir, output_dir = _ensure_dirs()

    output_ext = get_output_extension(conversion_type)
    output_filename = f"{task_uuid}{output_ext}"
    output_path = output_dir / output_filename

    staged_files: list[tuple[str, Path]] = []
    for index, (filename, payload) in enumerate(files, start=1):
        ext = Path(filename).suffix.lower() or ".bin"
        staged_path = upload_dir / f"{task_uuid}_{index:03d}{ext}"
        with open(staged_path, "wb") as f:
            f.write(payload)
        staged_files.append((filename, staged_path))

    import uuid as _uuid

    try:
        uid = _uuid.UUID(task_uuid)
    except (ValueError, AttributeError):
        uid = task_uuid

    task = db.query(ConversionTask).filter(ConversionTask.task_uuid == uid).first()
    if not task:
        logger.error("Task %s not found in DB", task_uuid)
        return

    task.status = TaskStatus.PROCESSING
    task.updated_at = datetime.now(timezone.utc)
    db.commit()

    try:
        runner = _resolve_batch_runner(conversion_type)
        if runner is None:
            raise RuntimeError(f"Неподдерживаемый пакетный тип конвертации: {conversion_type}")

        loop = asyncio.get_running_loop()
        await loop.run_in_executor(None, runner, staged_files, output_path, task_uuid)

        task.status = TaskStatus.DONE
        task.output_filename = output_filename
        task.updated_at = datetime.now(timezone.utc)
        db.commit()
        logger.info("Task %s completed successfully", task_uuid)

    except Exception as exc:
        logger.exception("Batch conversion failed for task %s: %s", task_uuid, exc)
        task.status = TaskStatus.FAILED
        task.error_message = str(exc)
        task.updated_at = datetime.now(timezone.utc)
        db.commit()

    finally:
        for _, staged_path in staged_files:
            try:
                os.remove(staged_path)
            except OSError:
                pass


def create_task_record(
    db: Session,
    user_id: Optional[int],
    original_filename: str,
    conversion_type: str,
) -> str:
    """Create a new ConversionTask row and return its UUID."""
    task = ConversionTask(
        task_uuid=uuid.uuid4(),
        user_id=user_id,
        original_filename=original_filename,
        conversion_type=conversion_type,
        status=TaskStatus.PENDING,
    )
    db.add(task)
    db.commit()
    db.refresh(task)
    return str(task.task_uuid)

